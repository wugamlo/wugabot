"""
WugaBot - Chat Interface Backend

This Flask application serves as the backend for the WugaBot chat interface.
It handles API requests, file processing, and AI model communication.

Author: WugaBot Team
Version: 1.0
"""

from flask import Flask, Response, render_template, request
from openai import OpenAI
import os
import json
import PyPDF2
import docx
import io
import logging

# Configure logging
log_level = os.getenv('LOG_LEVEL', 'INFO')
logging.basicConfig(
    level=getattr(logging, log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

import requests

@app.route('/models')
def get_models():
    """
    Retrieves available AI models from the Venice API

    Returns:
        JSON response with available models or error information
    """
    try:
        response = requests.get(
            "https://api.venice.ai/api/v1/models",
            headers={
                "Authorization": f"Bearer {os.getenv('VENICE_API_KEY')}",
                "Content-Type": "application/json"
            }
        )
        response.raise_for_status()
        models_data = response.json()
        models = []
        for model in models_data['data']:
            # Pass the entire model structure to the client
            # This includes model_spec with offline status and all capabilities
            models.append(model)
        return json.dumps({'models': models})
    except Exception as e:
        logger.error(f"Error fetching models: {str(e)}")
        return json.dumps({'error': str(e)}), 500

@app.route('/')
def index():
    """
    Serves the main application page

    Returns:
        Rendered HTML template
    """
    return render_template('index.html')



@app.route('/chat/expert', methods=['POST'])
def chat_expert():
    """
    Handles expert mode chat with multiple model candidates and synthesis
    
    Accepts:
        - JSON request with messages, candidate models, synthesis model, and parameters
        
    Returns:
        - JSON response with individual candidates and synthesized final answer
    """
    try:
        data = request.json
        messages = data.get('messages', [])
        candidate_models = data.get('candidate_models', [])
        synthesis_model = data.get('synthesis_model', 'mistral-31-24b')
        show_candidates = data.get('show_candidates', False)
        temperature = data.get('temperature', 0.7)
        max_completion_tokens = data.get('max_completion_tokens', 4000)
        candidate_capabilities = data.get('candidate_capabilities', {})
        synthesis_capabilities = data.get('synthesis_capabilities', {})
        
        logger.info(f"Expert mode request: {len(candidate_models)} candidates, synthesis: {synthesis_model}")
        logger.info(f"Candidate models: {candidate_models}")
        logger.info(f"Synthesis model from request: {synthesis_model}")
        
        if not candidate_models:
            return json.dumps({'error': 'No candidate models selected'}), 400
            
        # Generate responses from candidate models in parallel
        candidate_responses = []
        
        import concurrent.futures
        import threading
        
        def get_candidate_response(model):
            """Get response from a single candidate model"""
            try:
                venice_params = {
                    "include_venice_system_prompt": False
                }
                
                # Enable web search for web-capable models
                model_caps = candidate_capabilities.get(model, {})
                if model_caps.get('supportsWebSearch', False):
                    venice_params["enable_web_search"] = "on"
                    venice_params["enable_web_citations"] = True
                
                payload = {
                    "model": model,
                    "messages": messages,
                    "venice_parameters": venice_params,
                    "max_completion_tokens": max_completion_tokens,
                    "temperature": temperature,
                    "stream": False  # Non-streaming for candidates
                }
                
                response = requests.post(
                    "https://api.venice.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {os.getenv('VENICE_API_KEY')}",
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=60
                )
                
                if response.ok:
                    result = response.json()
                    if 'choices' in result and result['choices']:
                        content = result['choices'][0]['message']['content']
                        return {'model': model, 'content': content, 'success': True}
                    else:
                        return {'model': model, 'content': f"No response from {model}", 'success': False}
                else:
                    return {'model': model, 'content': f"Error from {model}: {response.status_code}", 'success': False}
                    
            except Exception as e:
                logger.error(f"Error getting response from {model}: {str(e)}")
                return {'model': model, 'content': f"Error: {str(e)}", 'success': False}
        
        # Execute candidate requests in parallel with improved error handling
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(candidate_models), 5)) as executor:
            future_to_model = {executor.submit(get_candidate_response, model): model for model in candidate_models}
            
            # Process completed futures with individual timeouts
            for future in concurrent.futures.as_completed(future_to_model, timeout=180):
                try:
                    result = future.result(timeout=60)  # Individual future timeout
                    candidate_responses.append(result)
                    logger.info(f"Received response from {result['model']}: success={result['success']}")
                except concurrent.futures.TimeoutError:
                    model = future_to_model[future]
                    logger.warning(f"Timeout for model {model}")
                    candidate_responses.append({
                        'model': model, 
                        'content': f"Timeout error for {model}", 
                        'success': False
                    })
                except Exception as e:
                    model = future_to_model[future]
                    logger.error(f"Error processing future for {model}: {str(e)}")
                    candidate_responses.append({
                        'model': model, 
                        'content': f"Processing error for {model}: {str(e)}", 
                        'success': False
                    })
        
        # Filter successful responses
        successful_responses = [r for r in candidate_responses if r['success']]
        
        if not successful_responses:
            return json.dumps({'error': 'All candidate models failed to respond'}), 500
        
        # Create synthesis prompt
        synthesis_messages = messages.copy()
        
        # Add candidate responses to synthesis prompt
        candidates_text = "\n\n".join([
            f"Response from {resp['model']}:\n{resp['content']}" 
            for resp in successful_responses
        ])
        
        synthesis_prompt = f"""You are tasked with synthesizing multiple AI responses into a single, comprehensive answer. Below are responses from different AI models to the same query.

Please create a synthesized response that:
1. Combines the best insights from all responses
2. Maintains consistency and coherence
3. Removes redundancy while preserving important details
4. Provides a balanced and well-structured answer

Candidate Responses:
{candidates_text}

Please provide a synthesized response that incorporates the strengths of each candidate while maintaining clarity and coherence."""

        synthesis_messages.append({'role': 'user', 'content': synthesis_prompt})
        
        # Get synthesis response with better error handling
        logger.info(f"Starting synthesis with model: {synthesis_model}")
        
        # Build venice parameters for synthesis
        synthesis_venice_params = {
            "include_venice_system_prompt": False
        }
        
        # Enable web search for synthesis model if it supports it
        synthesis_caps = synthesis_capabilities.get(synthesis_model, {})
        if synthesis_caps.get('supportsWebSearch', False):
            synthesis_venice_params["enable_web_search"] = "on"
            synthesis_venice_params["enable_web_citations"] = True
        
        synthesis_payload = {
            "model": synthesis_model,
            "messages": synthesis_messages,
            "venice_parameters": synthesis_venice_params,
            "max_completion_tokens": max_completion_tokens,
            "temperature": 0.3,  # Lower temperature for more consistent synthesis
            "stream": False
        }
        
        try:
            synthesis_response = requests.post(
                "https://api.venice.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {os.getenv('VENICE_API_KEY')}",
                    "Content-Type": "application/json"
                },
                json=synthesis_payload,
                timeout=120
            )
            
            if synthesis_response.ok:
                synthesis_result = synthesis_response.json()
                if 'choices' in synthesis_result and synthesis_result['choices']:
                    synthesized_content = synthesis_result['choices'][0]['message']['content']
                    logger.info("Synthesis completed successfully")
                else:
                    synthesized_content = "Failed to synthesize responses - no choices in response"
                    logger.error("Synthesis response missing choices")
            else:
                error_text = synthesis_response.text
                synthesized_content = f"Synthesis failed: {synthesis_response.status_code} - {error_text}"
                logger.error(f"Synthesis API error: {synthesis_response.status_code} - {error_text}")
        
        except requests.exceptions.Timeout:
            synthesized_content = f"Synthesis timed out using model {synthesis_model}"
            logger.error(f"Synthesis timeout with model: {synthesis_model}")
        except Exception as e:
            synthesized_content = f"Synthesis error: {str(e)}"
            logger.error(f"Synthesis exception: {str(e)}")
        
        # Prepare response
        response_data = {
            'synthesized_response': synthesized_content,
            'synthesis_model': synthesis_model,
            'candidate_count': len(successful_responses)
        }
        
        # Include individual candidates if requested
        if show_candidates:
            response_data['candidates'] = [
                {'model': resp['model'], 'content': resp['content']} 
                for resp in successful_responses
            ]
        
        return json.dumps(response_data), 200
        
    except Exception as e:
        logger.exception(f"Expert mode error: {str(e)}")
        return json.dumps({'error': f'Expert mode error: {str(e)}'}), 500

@app.route('/chat/stream', methods=['POST'])
def chat_stream():
    """
    Handles streaming chat completions from the AI model

    Accepts:
        - JSON request with messages, model selection, and parameters

    Returns:
        - Streaming response with AI-generated content
    """
    data = request.json
    search_enabled = data.get('web_search', False)
    messages = data.get('messages', [])

    # Use max_completion_tokens as the primary parameter, but fall back to max_tokens for backward compatibility
    max_completion_tokens = data.get('max_completion_tokens', data.get('max_tokens', 4000))
    temperature = data.get('temperature', 0.7)

    def generate(model, messages, temperature, max_completion_tokens, search_enabled):
        """
        Generator function that streams AI responses

        Args:
            model (str): The AI model to use
            messages (list): Chat history to send to the model
            temperature (float): Temperature parameter for generation
            max_completion_tokens (int): Maximum tokens to generate
            search_enabled (bool): Whether to enable web search

        Yields:
            Streaming response data from the AI model
        """
        try:
            logger.info(f"Generating response for model: {model}")
            logger.info(f"Web search setting: {search_enabled}")
            logger.info(f"Max completion tokens: {max_completion_tokens}")

            # Prepare the payload for Venice API with proper parameter names
            payload = {
                "model": model,
                "messages": messages,
                "venice_parameters": {
                    "include_venice_system_prompt": False
                },
                "max_completion_tokens": max_completion_tokens,  # Using specified parameter name
                "temperature": temperature,
                "stream": True
            }

            # Only add web search parameter when explicitly enabled
            if search_enabled == "on":
                payload["venice_parameters"]["enable_web_search"] = "on"
                payload["venice_parameters"]["enable_web_citations"] = True
                payload["venice_parameters"]["include_search_results_in_stream"] = True

            # Make request to Venice API
            logger.debug(f"Sending request to Venice API with payload: {json.dumps(payload)}")
            response = requests.post(
                "https://api.venice.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {os.getenv('VENICE_API_KEY')}",
                    "Content-Type": "application/json"
                },
                json=payload,
                stream=True
            )

            if not response.ok:
                logger.error(f"Venice API error: Status {response.status_code}")
                logger.error(f"Response content: {response.text}")
                yield f"data: {json.dumps({'error': f'API error: {response.status_code}'})}\n\n"
                return

            # Stream the response with improved handling
            for line in response.iter_lines():
                if not line:
                    continue

                line = line.decode('utf-8')
                if not line.startswith('data: '):
                    continue

                data = line[6:]
                if data == '[DONE]':
                    yield "data: [DONE]\n\n"
                    break

                try:
                    json_data = json.loads(data)

                    # Forward venice_parameters at the top level
                    if 'venice_parameters' in json_data:
                        # Handle citations separately to ensure proper JSON formatting
                        if 'web_search_citations' in json_data['venice_parameters']:
                            citations = json_data['venice_parameters']['web_search_citations']

                            if isinstance(citations, list) and len(citations) > 0:
                                # Clean and validate each citation before sending
                                cleaned_citations = []
                                for citation in citations:
                                    try:
                                        # Simplified citation with only title and URL
                                        cleaned_citation = {
                                            "title": str(citation.get("title", "")).strip() or "Untitled",
                                            "url": str(citation.get("url", "")).strip() or "#"
                                        }
                                        # Only add if we have at least a title or URL
                                        if cleaned_citation["title"] != "Untitled" or cleaned_citation["url"] != "#":
                                            cleaned_citations.append(cleaned_citation)
                                    except Exception as clean_error:
                                        logger.warning(f"Error cleaning citation: {clean_error}")
                                        continue

                                if cleaned_citations:
                                    # Send citations as a separate, well-formed JSON chunk
                                    try:
                                        citations_chunk = {
                                            "venice_parameters": {
                                                "web_search_citations": cleaned_citations
                                            }
                                        }
                                        # Use separators to ensure compact, clean JSON
                                        citations_json = json.dumps(citations_chunk, separators=(',', ':'), ensure_ascii=False)
                                        yield f"data: {citations_json}\n\n"
                                    except Exception as citation_error:
                                        logger.error(f"Error formatting citations: {citation_error}")

                        # Send other venice_parameters without citations to avoid duplication
                        other_params = {k: v for k, v in json_data['venice_parameters'].items() 
                                      if k != 'web_search_citations'}
                        if other_params:
                            other_chunk = {
                                "venice_parameters": other_params
                            }
                            yield f"data: {json.dumps(other_chunk)}\n\n"

                    # Process content and reasoning_content if present
                    if 'content' in json_data:
                        yield f"data: {json.dumps({'content': json_data['content']})}\n\n"

                    # Use standardized field for reasoning content
                    if 'reasoning_content' in json_data:
                        logger.debug(f"Found reasoning_content at top level: {json_data['reasoning_content'][:100]}...")
                        yield f"data: {json.dumps({'reasoning_content': json_data['reasoning_content']})}\n\n"

                    # Process delta content for streaming
                    if 'choices' in json_data and json_data['choices'] and 'delta' in json_data['choices'][0]:
                        delta = json_data['choices'][0]['delta']

                        # Stream content
                        if 'content' in delta and delta['content']:
                            yield f"data: {json.dumps({'content': delta['content']})}\n\n"

                        # Stream reasoning content
                        if 'reasoning_content' in delta and delta['reasoning_content']:
                            yield f"data: {json.dumps({'reasoning_content': delta['reasoning_content']})}\n\n"

                        # Stream venice parameters
                        if 'venice_parameters' in delta:
                            # Handle delta citations separately
                            if 'web_search_citations' in delta['venice_parameters']:
                                citations = delta['venice_parameters']['web_search_citations']

                                if isinstance(citations, list) and len(citations) > 0:
                                    # Clean and validate delta citations too
                                    cleaned_delta_citations = []
                                    for citation in citations:
                                        try:
                                            # Simplified citation with only title and URL
                                            cleaned_citation = {
                                                "title": str(citation.get("title", "")).strip() or "Untitled",
                                                "url": str(citation.get("url", "")).strip() or "#"
                                            }
                                            # Only add if we have at least a title or URL
                                            if cleaned_citation["title"] != "Untitled" or cleaned_citation["url"] != "#":
                                                cleaned_delta_citations.append(cleaned_citation)
                                        except Exception as clean_error:
                                            logger.warning(f"Error cleaning delta citation: {clean_error}")
                                            continue

                                    if cleaned_delta_citations:
                                        # Send delta citations as separate chunk
                                        try:
                                            delta_citations_chunk = {
                                                "venice_parameters": {
                                                    "web_search_citations": cleaned_delta_citations
                                                }
                                            }
                                            delta_citations_json = json.dumps(delta_citations_chunk, separators=(',', ':'), ensure_ascii=False)
                                            yield f"data: {delta_citations_json}\n\n"
                                        except Exception as delta_citation_error:
                                            logger.error(f"Error formatting delta citations: {delta_citation_error}")

                            # Send other delta venice_parameters without citations
                            other_delta_params = {k: v for k, v in delta['venice_parameters'].items() 
                                                if k != 'web_search_citations'}
                            if other_delta_params:
                                other_delta_chunk = json_data.copy()
                                other_delta_chunk['choices'][0]['delta']['venice_parameters'] = other_delta_params
                                yield f"data: {json.dumps(other_delta_chunk)}\n\n"

                except json.JSONDecodeError as e:
                    logger.warning(f"JSON decode error: {str(e)}, data: {data[:100]}...")
                    continue

        except Exception as e:
            logger.exception(f"Error in generate: {str(e)}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(
        generate(
            model=data.get('model', 'mistral-31-24b'),
            messages=data.get('messages', []),
            temperature=temperature,
            max_completion_tokens=max_completion_tokens,
            search_enabled=search_enabled
        ), 
        mimetype='text/event-stream'
    )

def extract_text_from_file(file_data, file_type):
    """
    Extracts text content from various file formats

    Supports txt, pdf, doc/docx, and xls/xlsx files

    Args:
        file_data (bytes): Binary file data
        file_type (str): File extension indicating the type

    Returns:
        str: Extracted text content or None if extraction fails
    """
    try:
        logger.info(f"Extracting text from {file_type} file")
        text = ""
        if file_type == 'txt':
            text = file_data.decode('utf-8')
            logger.debug("Text file decoded successfully")
        elif file_type == 'pdf':
            try:
                # Try using PyMuPDF first
                import fitz  # PyMuPDF
                pdf_file = io.BytesIO(file_data)
                pdf_document = fitz.open(stream=pdf_file, filetype="pdf")

                logger.debug(f"PDF file loaded, pages: {len(pdf_document)}")

                # Process each page
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]

                    # Extract text with better formatting preservation
                    page_text = page.get_text("text")

                    # Safer table extraction
                    try:
                        tables = page.find_tables()
                        if tables and hasattr(tables, 'tables') and tables.tables:
                            for table in tables.tables:
                                rows = []
                                for cells in table.rows:
                                    row_text = " | ".join([page.get_text("text", clip=cell.rect) for cell in cells])
                                    rows.append(row_text)
                                table_text = "\n".join(rows)
                                page_text += f"\n\n--- Table ---\n{table_text}\n--- End Table ---\n"
                    except Exception as table_err:
                        logger.warning(f"Table extraction error: {str(table_err)}")

                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"

                pdf_document.close()
            except Exception as fitz_err:
                # Fallback to PyPDF2 if PyMuPDF fails
                logger.warning(f"PyMuPDF failed: {str(fitz_err)}, falling back to PyPDF2")
                pdf_file = io.BytesIO(file_data)
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text() or "No text extracted"
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        elif file_type in ['doc', 'docx']:
            doc_file = io.BytesIO(file_data)
            doc = docx.Document(doc_file)
            logger.debug(f"DOC file loaded, paragraphs: {len(doc.paragraphs)}")
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_type in ['xls', 'xlsx']:
            import pandas as pd
            excel_file = io.BytesIO(file_data)
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)

            # Process each sheet
            tables = []
            for sheet_name, df in excel_data.items():
                # Convert DataFrame to string representation with proper formatting
                table_text = f"\n--- Sheet: {sheet_name} ---\n"
                table_text += df.to_string(index=False)
                tables.append(table_text)

            text = "\n\n".join(tables)
            logger.debug(f"Excel file processed, found {len(excel_data)} sheets")

        if not text:
            logger.warning("Warning: Extracted text is empty")
            return None

        logger.info(f"Successfully extracted {len(text)} characters")
        return text.strip()
    except Exception as e:
        logger.exception(f"Error extracting text: {str(e)}")
        return None

@app.route('/process_file', methods=['POST'])
def process_file():
    """
    Processes uploaded files and extracts their text content

    Handles various file formats including txt, pdf, doc/docx, and xls/xlsx
    Enforces file size limits and provides appropriate error messages

    Returns:
        JSON response with extracted text or error information
    """
    try:
        logger.info("Starting file processing...")

        if 'file' not in request.files:
            logger.warning("No file in request.files")
            return json.dumps({'error': 'No file part'}), 400, {'Content-Type': 'application/json'}

        file = request.files['file']
        logger.info(f"Received file: {file.filename}")

        if file.filename == '':
            logger.warning("Empty filename received")
            return json.dumps({'error': 'No file selected'}), 400, {'Content-Type': 'application/json'}

        # Set CORS headers
        headers = {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST',
            'Access-Control-Allow-Headers': 'Content-Type'
        }

        # Check file size (2MB limit)
        file_data = file.read()
        logger.debug(f"File size: {len(file_data)} bytes")
        if len(file_data) > 2 * 1024 * 1024:  # 2MB in bytes
            return json.dumps({'error': 'File too large. Maximum size is 2MB'}), 400, {'Content-Type': 'application/json'}

        file_type = file.filename.split('.')[-1].lower()
        if file_type not in ['txt', 'pdf', 'doc', 'docx', 'xls', 'xlsx']:
            return json.dumps({'error': f'Unsupported file type: {file_type}'}, ensure_ascii=False), 400, {'Content-Type': 'application/json'}

        extracted_text = extract_text_from_file(file_data, file_type)
        if extracted_text is None:
            return json.dumps({'error': 'Failed to extract text from file'}, ensure_ascii=False), 400

        return json.dumps({'text': extracted_text}, ensure_ascii=False), 200, headers
    except Exception as e:
        logger.exception(f"File processing error: {str(e)}")
        return json.dumps({'error': f'File processing error: {str(e)}'}, ensure_ascii=False), 500, headers

@app.route('/generate_visualization', methods=['POST'])
def generate_visualization():
    """
    Generates visualizations based on the request type and data

    Supports charts, diagrams, and simple drawings

    Args:
        Request JSON with visualization_type and data fields

    Returns:
        JSON with the generated visualization as SVG or base64 image
    """
    try:
        # Make sure we have a valid JSON request
        if not request.is_json:
            logger.error("Invalid request: Not JSON")
            return json.dumps({'error': 'Invalid request format. Expected JSON.'}), 400

        data = request.json
        logger.info(f"Visualization request received: {data.keys()}")

        visualization_type = data.get('visualization_type')
        if not visualization_type:
            logger.error("Missing visualization_type in request")
            return json.dumps({'error': 'Missing visualization_type parameter'}), 400

        # Validate visualization type is one of the supported types
        if visualization_type not in ['chart', 'diagram', 'drawing']:
            logger.error(f"Unsupported visualization type: {visualization_type}")
            return json.dumps({'error': f'Unsupported visualization type: {visualization_type}'}), 400

        # Simplified handling of visualization data
        viz_data = data.get('data', {})
        logger.info(f"Raw visualization data type: {type(viz_data)}")

        # Ensure we have a valid data object with proper defaults
        if not viz_data or not isinstance(viz_data, dict):
            logger.warning("Invalid visualization data received, using defaults")
            if visualization_type == 'chart':
                viz_data = {
                    "chart_type": "bar",
                    "title": "Default Chart",
                    "labels": ["A", "B", "C"],
                    "values": [10, 20, 30]
                }
            elif visualization_type == 'diagram':
                viz_data = {
                    "diagram_type": "flowchart",
                    "elements": [
                        {"text": "Start"},
                        {"text": "Process"},
                        {"text": "End"}
                    ]
                }
            elif visualization_type == 'drawing':
                viz_data = {
                    "description": "A simple drawing"
                }
            else:
                viz_data = {"default": True}

            logger.info(f"Using default data for {visualization_type}")

        # Create a deep copy to avoid modifying the original
        sanitized_data = {}

        try:
            # Validate key structure and fill in missing fields with defaults
            if visualization_type == 'chart':
                # Ensure required fields exist with proper defaults
                sanitized_data["chart_type"] = viz_data.get("chart_type", "bar")
                if not isinstance(sanitized_data["chart_type"], str):
                    sanitized_data["chart_type"] = "bar"

                sanitized_data["title"] = viz_data.get("title", "Chart")
                if not isinstance(sanitized_data["title"], str):
                    sanitized_data["title"] = "Chart"

                sanitized_data["labels"] = viz_data.get("labels", ["A", "B", "C"])
                if not isinstance(sanitized_data["labels"], list):
                    sanitized_data["labels"] = ["A", "B", "C"]

                sanitized_data["values"] = viz_data.get("values", [10, 20, 30])
                if not isinstance(sanitized_data["values"], list):
                    sanitized_data["values"] = [10, 20, 30]

                # Ensure values are numeric with stronger validation
                sanitized_data["values"] = []
                for v in viz_data.get("values", [10, 20, 30]):
                    try:
                        if isinstance(v, (int, float)):
                            sanitized_data["values"].append(float(v))
                        elif isinstance(v, str) and v.replace('.', '', 1).isdigit():
                            sanitized_data["values"].append(float(v))
                        else:
                            sanitized_data["values"].append(0)
                    except:
                        sanitized_data["values"].append(0)

                # Make sure we have at least some data
                if not sanitized_data["values"] or len(sanitized_data["values"]) == 0:
                    sanitized_data["values"] = [10, 20, 30]

                # Use sanitized data
                viz_data = sanitized_data
                logger.info(f"Validated chart data: {viz_data}")

            elif visualization_type == 'diagram':
                sanitized_data["diagram_type"] = viz_data.get("diagram_type", "flowchart")
                if not isinstance(sanitized_data["diagram_type"], str):
                    sanitized_data["diagram_type"] = "flowchart"

                sanitized_data["elements"] = []
                # Validate each element
                for elem in viz_data.get("elements", [{"text": "Start"}, {"text": "Process"}, {"text": "End"}]):
                    if isinstance(elem, dict) and "text" in elem:
                        sanitized_data["elements"].append({"text": str(elem["text"])})
                    else:
                        # Skip invalid elements
                        logger.warning(f"Skipping invalid diagram element: {elem}")

                # If no valid elements, use defaults
                if not sanitized_data["elements"]:
                    sanitized_data["elements"] = [{"text": "Start"}, {"text": "Process"}, {"text": "End"}]

                # Use sanitized data
                viz_data = sanitized_data

            elif visualization_type == 'drawing':
                sanitized_data["description"] = str(viz_data.get("description", "A simple drawing"))
                # Use sanitized data
                viz_data = sanitized_data

        except Exception as validation_error:
            logger.error(f"Error during data validation: {str(validation_error)}")
            # Fall back to safe defaults
            if visualization_type == 'chart':
                viz_data = {
                    "chart_type": "bar",
                    "title": "Default Chart",
                    "labels": ["A", "B", "C"],
                    "values": [10, 20, 30]
                }
            elif visualization_type == 'diagram':
                viz_data = {
                    "diagram_type": "flowchart",
                    "elements": [
                        {"text": "Start"},
                        {"text": "Process"},
                        {"text": "End"}
                    ]
                }
            elif visualization_type == 'drawing':
                viz_data = {
                    "description": "A simple drawing"
                }

        # Log the final data being used
        logger.info(f"Using data for visualization: {str(viz_data)[:200]}")

        if visualization_type == 'chart':
            # Generate chart using matplotlib
            try:
                import matplotlib.pyplot as plt
                import matplotlib
                matplotlib.use('Agg')
                import io
                import base64

                chart_type = viz_data.get('chart_type', 'bar')
                title = viz_data.get('title', 'Chart')
                labels = viz_data.get('labels', [])
                values = viz_data.get('values', [])

                # Validate input data
                if not labels or not values:
                    logger.warning(f"Missing data for chart generation - Labels: {labels}, Values: {values}")
                    # Use default data if missing
                    if not labels:
                        labels = ["No Data"] if not values else [f"Item {i+1}" for i in range(len(values))]
                    if not values:
                        values = [0] if not labels else [10 for _ in range(len(labels))]

                # Ensure values are numeric
                values = [float(v) if isinstance(v, (int, float, str)) else 0 for v in values]

                plt.figure(figsize=(10, 6))

                logger.info(f"Generating {chart_type} chart with {len(labels)} labels and {len(values)} values")

                if chart_type == 'bar':
                    plt.bar(labels, values)
                elif chart_type == 'line':
                    plt.plot(labels, values)
                elif chart_type == 'pie':
                    # Ensure no negative values for pie charts
                    pie_values = [max(0, v) for v in values]
                    # If all values are 0, use default values
                    if sum(pie_values) == 0:
                        pie_values = [1 for _ in range(len(labels))]
                    plt.pie(pie_values, labels=labels, autopct='%1.1f%%')
                else:
                    # Default to bar chart for unknown types
                    logger.warning(f"Unknown chart type: {chart_type}, defaulting to bar")
                    plt.bar(labels, values)

                plt.title(title)
                plt.tight_layout()

                # Convert plot to base64 image
                buf = io.BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                img_str = base64.b64encode(buf.read()).decode('utf-8')
                plt.close()
            except Exception as chart_error:
                logger.exception(f"Error generating chart: {chart_error}")
                # Return a simple error chart
                import numpy as np
                plt.figure(figsize=(10, 6))
                plt.text(0.5, 0.5, f"Error generating chart: {str(chart_error)}", 
                        horizontalalignment='center', verticalalignment='center',
                        transform=plt.gca().transAxes, color='red')
                plt.tight_layout()
                buf = io.BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                img_str = base64.b64encode(buf.read()).decode('utf-8')
                plt.close()

            return json.dumps({
                'image': f'data:image/png;base64,{img_str}',
                'type': 'chart'
            })

        elif visualization_type == 'diagram':
            # Generate SVG diagram
            try:
                # First make sure svgwrite is installed
                try:
                    import svgwrite
                except ImportError:
                    logger.error("svgwrite module not found, installing...")
                    import subprocess
                    subprocess.check_call(["pip", "install", "svgwrite"])
                    import svgwrite

                diagram_type = viz_data.get('diagram_type', 'flowchart')
                elements = viz_data.get('elements', [])

                logger.info(f"Generating {diagram_type} diagram with {len(elements) if elements else 0} elements")

                # Create a simple SVG drawing with white background for visibility
                dwg = svgwrite.Drawing('diagram.svg', profile='tiny', size=('800px', '600px'))
                # Add a background rectangle
                dwg.add(dwg.rect((0, 0), ('100%', '100%'), fill='#ffffff'))
            except Exception as diagram_setup_error:
                logger.exception(f"Error setting up diagram: {diagram_setup_error}")
                return json.dumps({'error': f'Error setting up diagram: {str(diagram_setup_error)}'}), 500

            # Default elements for a simple flowchart if none provided
            if not elements and diagram_type == 'flowchart':
                elements = [
                    {'text': 'Start'},
                    {'text': 'Decision\n(Yes or No?)'},
                    {'text': 'Yes', 'branch': 'left'},
                    {'text': 'No', 'branch': 'right'},
                    {'text': 'End', 'branch': 'left'},
                    {'text': 'End', 'branch': 'right'}
                ]

            if diagram_type == 'flowchart':
                # Improved flowchart implementation
                main_y_pos = 50

                # Draw the start node
                dwg.add(dwg.rect((325, main_y_pos), (150, 80), fill='#f0f0f0', stroke='#000000', rx=5, ry=5))
                dwg.add(dwg.text(elements[0].get('text', 'Start'), insert=(400, main_y_pos + 45), 
                                text_anchor="middle", font_size=16))

                # Draw connector
                main_y_pos += 100
                dwg.add(dwg.line((400, main_y_pos - 20), (400, main_y_pos + 20), stroke='#000000', stroke_width=2))
                dwg.add(dwg.polygon([(395, main_y_pos + 10), (400, main_y_pos + 20), (405, main_y_pos + 10)], 
                                  fill='#000000'))

                # Draw decision node
                main_y_pos += 50
                dwg.add(dwg.polygon([(400, main_y_pos), (475, main_y_pos + 50), (400, main_y_pos + 100), (325, main_y_pos + 50)], 
                                  fill='#f0f0f0', stroke='#000000'))

                # Multi-line text for decision
                decision_text = elements[1].get('text', 'Decision?').split("\n")
                for i, line in enumerate(decision_text):
                    y_offset = main_y_pos + 50 + (i - len(decision_text)/2) * 20
                    dwg.add(dwg.text(line, insert=(400, y_offset), text_anchor="middle", font_size=14))

                # Draw Yes/No paths
                main_y_pos += 120

                # Yes branch (left)
                dwg.add(dwg.line((350, main_y_pos - 20), (250, main_y_pos + 50), stroke='#000000', stroke_width=2))
                dwg.add(dwg.polygon([(260, main_y_pos + 45), (250, main_y_pos + 50), (255, main_y_pos + 35)], 
                                  fill='#000000'))
                dwg.add(dwg.text("Yes", insert=(290, main_y_pos + 10), text_anchor="middle", font_size=14))

                # No branch (right)
                dwg.add(dwg.line((450, main_y_pos - 20), (550, main_y_pos + 50), stroke='#000000', stroke_width=2))
                dwg.add(dwg.polygon([(545, main_y_pos + 45), (550, main_y_pos + 50), (540, main_y_pos + 40)], 
                                  fill='#000000'))
                dwg.add(dwg.text("No", insert=(510, main_y_pos + 10), text_anchor="middle", font_size=14))

                # Left node (Yes result)
                left_y = main_y_pos + 70
                dwg.add(dwg.rect((175, left_y), (150, 80), fill='#f0f0f0', stroke='#000000', rx=5, ry=5))
                dwg.add(dwg.text(elements[2].get('text', 'Yes'), insert=(250, left_y + 45), 
                                text_anchor="middle", font_size=16))

                # Right node (No result)
                dwg.add(dwg.rect((475, left_y), (150, 80), fill='#f0f0f0', stroke='#000000', rx=5, ry=5))
                dwg.add(dwg.text(elements[3].get('text', 'No'), insert=(550, left_y + 45), 
                                text_anchor="middle", font_size=16))

                # Draw connectors to End nodes
                left_y += 100

                # Left End connector
                dwg.add(dwg.line((250, left_y - 20), (250, left_y + 20), stroke='#000000', stroke_width=2))
                dwg.add(dwg.polygon([(245, left_y + 10), (250, left_y + 20), (255, left_y + 10)], 
                                  fill='#000000'))

                # Right End connector
                dwg.add(dwg.line((550, left_y - 20), (550, left_y + 20), stroke='#000000', stroke_width=2))
                dwg.add(dwg.polygon([(545, left_y + 10), (550, left_y + 20), (555, left_y + 10)], 
                                  fill='#000000'))

                # Left End node
                dwg.add(dwg.rect((175, left_y + 40), (150, 80), fill='#f0f0f0', stroke='#000000', rx=5, ry=5))
                dwg.add(dwg.text(elements[4].get('text', 'End'), insert=(250, left_y + 85), 
                                text_anchor="middle", font_size=16))

                # Right End node
                dwg.add(dwg.rect((475, left_y + 40), (150, 80), fill='#f0f0f0', stroke='#000000', rx=5, ry=5))
                dwg.add(dwg.text(elements[5].get('text', 'End'), insert=(550, left_y + 85), 
                                text_anchor="middle", font_size=16))

            svg_string = dwg.tostring()
            return json.dumps({
                'svg': svg_string,
                'type': 'diagram'
            })

        elif visualization_type == 'drawing':
            # Generate a simple drawing based on text description
            from PIL import Image, ImageDraw, ImageFont
            import io
            import base64

            description = viz_data.get('description', '')
            logger.info(f"Drawing description: {description}")

            # Create a blank canvas
            img = Image.new('RGB', (500, 500), color='white')
            draw = ImageDraw.Draw(img)

            # More detailed drawing based on keywords
            if 'cat' in description.lower():
                # Draw cat face
                draw.ellipse((100, 100, 400, 400), outline='black', width=3)  # Face

                # Draw cat ears
                draw.polygon([(150, 150), (200, 50), (250, 150)], fill='white', outline='black', width=3)  # Left ear
                draw.polygon([(350, 150), (300, 50), (250, 150)], fill='white', outline='black', width=3)  # Right ear

                # Draw cat eyes
                draw.ellipse((175, 200, 225, 250), fill='white', outline='black', width=2)  # Left eye
                draw.ellipse((275, 200, 325, 250), fill='white', outline='black', width=2)  # Right eye

                # Draw pupils
                draw.ellipse((190, 215, 210, 235), fill='black')  # Left pupil
                draw.ellipse((290, 215, 310, 235), fill='black')  # Right pupil

                # Draw nose
                draw.polygon([(250, 270), (230, 290), (270, 290)], fill='pink', outline='black')

                # Draw whiskers
                for i in range(3):
                    # Left whiskers
                    draw.line((170, 290 + i*15, 70, 270 + i*15), fill='black', width=2)
                    # Right whiskers
                    draw.line((330, 290 + i*15, 430, 270 + i*15), fill='black', width=2)

                # Draw smile
                draw.arc((200, 280, 300, 350), 0, 180, fill='black', width=3)

            elif 'circle' in description.lower():
                draw.ellipse((100, 100, 400, 400), outline='black', width=3, fill='#FFEEEE')
            elif 'square' in description.lower():
                draw.rectangle((100, 100, 400, 400), outline='black', width=3, fill='#EEEEFF')
            elif 'triangle' in description.lower():
                draw.polygon([(250, 100), (100, 400), (400, 400)], outline='black', width=3, fill='#EEFFEE')
            elif 'smiley' in description.lower() or 'face' in description.lower():
                draw.ellipse((100, 100, 400, 400), outline='black', width=3, fill='#FFFFEE')  # Face
                draw.ellipse((160, 180, 210, 230), fill='black')  # Left eye
                draw.ellipse((290, 180, 340, 230), fill='black')  # Right eye
                draw.arc((150, 200, 350, 350), 0, 180, fill='black', width=5)  # Smile
            else:
                # Default drawing - simple cartoon character
                draw.ellipse((150, 100, 350, 300), outline='black', width=3, fill='#FFFFEE')  # Head
                draw.ellipse((200, 150, 230, 180), fill='black')  # Left eye
                draw.ellipse((270, 150, 300, 180), fill='black')  # Right eye
                draw.arc((200, 200, 300, 250), 0, 180, fill='black', width=3)  # Smile

                # Add a label with the description
                try:
                    # Try to load a font, but don't fail if not available
                    try:
                        font = ImageFont.truetype("Arial", 20)
                    except:
                        # Fall back to default font
                        font = ImageFont.load_default()

                    # Add description text at the bottom
                    draw.text((250, 450), description, fill="black", anchor="ms", font=font)
                except Exception as font_error:
                    logger.warning(f"Font error: {str(font_error)}")

            # Convert to base64
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            buf.seek(0)
            img_str = base64.b64encode(buf.read()).decode('utf-8')

            return json.dumps({
                'image': f'data:image/png;base64,{img_str}',
                'type': 'drawing'
            })

        else:
            return json.dumps({'error': 'Unsupported visualization type'}), 400

    except Exception as e:
        logger.exception(f"Visualization generation error: {str(e)}")
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        import io
        import base64

        # Create a cleaner error visualization
        plt.figure(figsize=(8, 4), facecolor='white')
        plt.text(0.5, 0.7, "Visualization Error", 
                horizontalalignment='center', verticalalignment='center',
                transform=plt.gca().transAxes, color='#e74c3c', fontsize=16, fontweight='bold')

        # Simplified error message for display
        error_msg = str(e)
        if len(error_msg) > 100:
            error_msg = error_msg[:100] + "..."

        plt.text(0.5, 0.5, f"{error_msg}", 
                horizontalalignment='center', verticalalignment='center',
                transform=plt.gca().transAxes, color='#2c3e50', fontsize=12)

        plt.text(0.5, 0.3, f"Type: {visualization_type}", 
                horizontalalignment='center', verticalalignment='center',
                transform=plt.gca().transAxes, color='#7f8c8d', fontsize=10)

        # Remove axes for cleaner look
        plt.axis('off')
        plt.tight_layout()

        # Convert to base64
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        plt.close()

        # Return the error image with a cleaner response
        return json.dumps({
            'image': f'data:image/png;base64,{img_str}',
            'type': 'error',
            'error': str(e)
        })


# Helper function to import traceback module
def import_traceback():
    try:
        import traceback
        return traceback
    except ImportError:
        # Define a minimal fallback if traceback is not available
        class MinimalTraceback:
            def format_exc(self):
                return "Traceback information not available"
        return MinimalTraceback()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)